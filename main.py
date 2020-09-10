from tkinter import *
import tkinter as tk
import processes as ps
import time
import speech_recognition as sr
import docx
import os

root = tk.Tk()
root.title("SIPENILAI")
root.wm_iconbitmap('speechtotext.ico')

startTime = time.time()
n = 2
w = 2
p = 2
student = 1
question = 5
jmlquest = 5
#human_rater = [500]
#human_rater = [35, 73, 36, 62, 42, 75, 90, 65, 86, 68,
# 97, 57, 15, 77, 92, 97, 93, 96, 80, 87,
# 88, 98, 94, 96, 98, 79, 79, 81, 70, 74,
# 37, 81, 37, 77, 90, 73, 42, 96, 87, 94,
# 88, 63, 72]
no1 = []
excel = []
list_score = []
doc = docx.Document('soal1.docx')
test = []
arr_score = []
txt = []

class Sipenilai(tk.Tk):
    def __init__(self, *args, **kwargs):
        tk.Tk.__init__(self, *args, **kwargs)
        container = tk.Frame(root, bg='#80c1ff', bd=5)
        container.place(relwidth=0.75, relheight=0.1, relx=0.5, rely=0.1, anchor='n')
        container.pack(side="top", fill="both", expand=True)
        container.grid_rowconfigure(0, weight=1)
        container.grid_columnconfigure(0, weight=1)

        self.frames = {}

        for F in (StartPage, PageOne, PageTwo, PageThree, PageFour, PageFive):
            frame = F(container, self)

            self.frames[F] = frame

            frame.grid(row=0, column=0, sticky="nsew")

        self.show_frame(StartPage)

    def show_frame(self, cont):
        frame = self.frames[cont]
        frame.tkraise()

class StartPage(tk.Frame):

    def __init__(self, parent, controller):
        self.controller = controller
        tk.Frame.__init__(self,parent)

        label = tk.Label(self, text="Tekan next untuk memulai", font=40)
        label.pack(pady=10,padx=10)

        button = tk.Button(self, text="Next",
                            command=lambda: controller.show_frame(PageOne))
        button.pack()


class PageOne(tk.Frame):

    def __init__(self, parent, controller):
        self.controller = controller
        tk.Frame.__init__(self, parent)
        label = tk.Label(self, text=str(doc.paragraphs[0].text), font=40)
        label.pack(pady=10,padx=10)

        def rec():
            sample_rate = 48000
            chunk_size = 2048
            r = sr.Recognizer()

            with sr.Microphone(sample_rate = sample_rate, chunk_size = chunk_size) as source:
                r.adjust_for_ambient_noise(source)
                audio = r.listen(source)
            try:
                speak = r.recognize_google(audio, language="ja-JP")
                msg.configure(text=speak)
                print(speak)
                txt.insert(0, speak)
            except Exception as e:
                print(e)


        button2 = tk.Button(self, text="next",
                            command=lambda: controller.show_frame(PageTwo))
        button2.pack()
        msg = tk.Label(self, font=40)
        msg.pack(pady=20,padx=20)
        btn = tk.Button(self, text="Tekan untuk Menjawab", command=rec)
        btn.pack()


class PageTwo(tk.Frame):

    def __init__(self, parent, controller):
        self.controller = controller
        tk.Frame.__init__(self, parent)
        label = tk.Label(self, text=str(doc.paragraphs[1].text), font=40)
        label.pack(pady=10,padx=10)

        def rec():
            sample_rate = 48000
            chunk_size = 2048
            r = sr.Recognizer()

            with sr.Microphone(sample_rate = sample_rate, chunk_size = chunk_size) as source:
                r.adjust_for_ambient_noise(source)
                audio = r.listen(source)
            try:
                speak = r.recognize_google(audio, language="ja-JP")
                msg.configure(text=speak)
                print(speak)
                txt.insert(1, speak)

            except Exception as e:
                print(e)


        button1 = tk.Button(self, text="Next",
                            command=lambda: controller.show_frame(PageThree))
        button1.pack()


        msg = tk.Label(self, font=40)
        msg.pack(pady=20, padx=20)
        btn = tk.Button(self, text="Tekan untuk Menjawab", command=rec)
        btn.pack()

class PageThree(tk.Frame):

    def __init__(self, parent, controller):
        self.controller = controller
        tk.Frame.__init__(self, parent)
        label = tk.Label(self, text=str(doc.paragraphs[2].text), font=40)
        label.pack(pady=10,padx=10)
        def rec():
            sample_rate = 48000
            chunk_size = 2048
            r = sr.Recognizer()

            with sr.Microphone(sample_rate = sample_rate, chunk_size = chunk_size) as source:
                r.adjust_for_ambient_noise(source)
                audio = r.listen(source)
            try:
                speak = r.recognize_google(audio, language="ja-JP")
                msg.configure(text=speak)
                print(speak)
                txt.insert(2, speak)
            except Exception as e:
                print(e)


        button1 = tk.Button(self, text="Next",
                            command=lambda: controller.show_frame(PageFour))
        button1.pack()


        msg = tk.Label(self, font=40)
        msg.pack(pady=20, padx=20)
        btn = tk.Button(self, text="Tekan untuk Menjawab", command=rec)
        btn.pack()

class PageFour(tk.Frame):

    def __init__(self, parent, controller):
        self.controller = controller
        tk.Frame.__init__(self, parent)
        label = tk.Label(self, text=str(doc.paragraphs[3].text), font=40)
        label.pack(pady=10,padx=10)
        def rec():
            sample_rate = 48000
            chunk_size = 2048
            r = sr.Recognizer()

            with sr.Microphone(sample_rate = sample_rate, chunk_size = chunk_size) as source:
                r.adjust_for_ambient_noise(source)
                audio = r.listen(source)
            try:
                speak = r.recognize_google(audio, language="ja-JP")
                msg.configure(text=speak)
                print(speak)
                txt.insert(3, speak)
            except Exception as e:
                print(e)


        button1 = tk.Button(self, text="Next",
                            command=lambda: controller.show_frame(PageFive))
        button1.pack()


        msg = tk.Label(self, font=40)
        msg.pack(pady=20, padx=20)
        btn = tk.Button(self, text="Tekan untuk Menjawab", command=rec)
        btn.pack()

class PageFive(tk.Frame):

    def __init__(self, parent, controller):
        self.controller = controller
        tk.Frame.__init__(self, parent)
        label = tk.Label(self, text=str(doc.paragraphs[4].text), font=40)
        label.pack(pady=10,padx=10)

        def rec():
            sample_rate = 48000
            chunk_size = 2048
            r = sr.Recognizer()

            with sr.Microphone(sample_rate = sample_rate, chunk_size = chunk_size) as source:
                r.adjust_for_ambient_noise(source)
                audio = r.listen(source)
            try:
                speak = r.recognize_google(audio, language="ja-JP")
                msg.configure(text=speak)
                print(speak)
                txt.insert(4, speak)
            except Exception as e:
                print(e)

        msg = tk.Label(self, font=40)
        msg.pack(pady=20, padx=20)
        btn = tk.Button(self, text="Tekan untuk Menjawab", command=rec)
        btn.pack()

        def score():
            arr_score = []
            scrollbar = Scrollbar(root)
            scrollbar.pack(side=RIGHT, fill=Y)
            mylist = Listbox(root, yscrollcommand=scrollbar.set)

            for q in range(0, question):
                kj = 1
                scores = []

                # process the student's answer document
                prep = ps.preprocessing(txt[q])
                mylist.insert(END, '\nJAWABAN ' + str(q + 1), )
                mylist.pack(side=LEFT, fill=BOTH, expand=True)
                scrollbar.config(command=mylist.yview)

                print('\nJAWABAN ', q + 1)

                key = ps.read_txt("jwbDosen" + str(q + 1) + ".docx")  # read answer key documents
                # for each answer keys (from each questions)
                for x in range(0, len(key)):
                    # process answer keys
                    prep2 = ps.preprocessing(key[x])  # delete repetitive from question, convert to romaji(if needed), filter text
                    tdmkj = ps.TDMRef(prep2)
                    tdms = ps.TDMTest(prep2, prep)
                    svdkj = ps.SVD(tdmkj)
                    svds = ps.SVD(tdms)
                    frobnorm = ps.frobeniusNorm(svdkj, svds)
                    scores.append(frobnorm)

                    # #print all processes' results
                    print('----------------')
                    mylist.insert(END, 'KUNCI JAWABAN  ' + str(kj))
                    mylist.insert(END, '--\npreprocessing dosen\n  ' + str(prep2))
                    mylist.insert(END, '--\npreprocessing siswa\n  ' + str(prep))
                    print('KUNCI JAWABAN '+str(kj))
                    print('---\npreprocessing dosen\n', prep2)
                    print('---\ntdm kj\n', tdmkj)
                    print('---\npreprocessing siswa\n', prep)
                    print('---\ntdm student\n', tdms)
                    print('---\nSVD KJ: ', svdkj, "SVD S: ", svds, "Nilai KJ ke-", kj, ": ", frobnorm)
                    #print('---\nSVD S: ', svds)
                    print('---\nNilai KJ ke-', kj, '\t: ', frobnorm)

                    kj += 1  # loop for each answer keys
                # # print score of each questions (maximum score from list of each key answers' score)
                # print('\n\nNilai Soal Nomor', q + 1,'\t: ', max(scores), '\n--------------------------')
                # add the score of each questions to list of scores of each students
                arr_score.append(max(scores))
                # excel.append(max(scores))
                # # print students' score (sum of scores from each questions)
                # print('\n\nNILAI MAHASISWA ', count, '\t: ', round(sum(arr_score),2))
            list_score.append(round(sum(arr_score), 2))
            #akurasi = []
            #for n in range(student):
            #    acc = round((100 - (((abs(list_score[n] - human_rater[n])) / 100) * 100)), 2)
            #    akurasi.append(acc)
            mylist.insert(END, '=================================================================================== ' )
            #mylist.insert(END, 'Human Rater\t\t\t: ' + str(human_rater))
            mylist.insert(END, 'Nilai-nilai Siswa\t: ' + str(list_score))
            mylist.insert(END, 'Nilai per soal \t\t: ' + str(arr_score))
            mylist.insert(END, 'Program Execution Duration: ' + str(time.time() - startTime), "seconds")
            #mylist.insert(END, 'Akurasi\t\t\t\t: ' + str(akurasi))
            #print("Human Rater\t\t\t: ", human_rater)
            print("Nilai-nilai Siswa\t: ", list_score)
            print("Nilai per soal \t\t: ", arr_score)
            # print("Akurasi\t\t\t\t: ", akurasi)
            # print("Rata-rata Akurasi\t: ", round(statistics.mean(akurasi),2))
            # print("Standar Deviasi\t\t: ", round(statistics.stdev(akurasi),2))
            print("Program Execution Duration: ", (time.time() - startTime), "seconds")

        button1 = tk.Button(self, text="Submit", command=score)
        button1.pack()
        
       
app = Sipenilai()
app.mainloop()

'''
#create excel file of the data
wb = Workbook()                                             # Workbook is created
#wb = load_workbook('data.xlsx')
sheet1 = wb.add_sheet("Hasil_Test")                  # add_sheet is used to create sheet.
#for n in range(1, 43):
# for o in range(len(excel)):
#     for m in range(2,6):
#         sheet1.cell(row=n, column=1, value=n)
#         sheet1.cell(row=n, column=m, value=excel[o])
#for n in range(len(human_rater)):
# sheet1.cell(row=1, column=1, value='No')
# sheet1.cell(row=1, column=2, value='Jawaban')
# sheet1.cell(row=1, column=3, value='Total Score')
# sheet1.cell(row=1, column=4, value='Score')
# sheet1.cell(row=n+2, column=1, value=n)
sheet1.write(row=n+2, column=2, value=human_rater[n])
sheet1.write(row=n+2, column=3, value=list_score[n])
sheet1.write(row=n+2, column=4, value=arr_score[n])
#sheet1.cell(row=n+2, column=5, value=akurasi[n])
wb.save('data.xlsx')
'''
